"""Build an editable Word report from the filled IDC LaTeX report.

The converter intentionally supports the LaTeX vocabulary used by
``outputs/report/template-latex/template.tex``.  It is not a general-purpose
LaTeX converter: keeping the supported surface small makes the monthly DOCX
deterministic and lets its content remain aligned with the publication PDF.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.image.image import Image as DocxImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


BODY_FONT = "Palatino"
SANS_FONT = "Helvetica"
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(95, 95, 95)
PLACEHOLDER_RED = RGBColor(155, 35, 35)
FIGURE_WIDTH_MM = 150.0
FIGURE_MAX_HEIGHT_MM = {1: 150.0, 2: 205.0}
FIGURE_PLACEHOLDER_HEIGHT_MM = {1: 140.0, 2: 170.0}


def _extract_group(source: str, opening_brace: int) -> tuple[str, int]:
    if opening_brace >= len(source) or source[opening_brace] != "{":
        raise ValueError("Expected an opening brace")
    depth = 0
    escaped = False
    for index in range(opening_brace, len(source)):
        char = source[index]
        if escaped:
            escaped = False
            continue
        if char == "\\":
            escaped = True
            continue
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return source[opening_brace + 1:index], index + 1
    raise ValueError("Unbalanced LaTeX braces")


def _command_argument(source: str, command: str, start: int = 0) -> tuple[str, int, int]:
    match = re.search(rf"\\{re.escape(command)}\s*\{{", source[start:])
    if not match:
        raise ValueError(f"LaTeX command not found: \\{command}")
    command_start = start + match.start()
    brace = start + match.end() - 1
    value, end = _extract_group(source, brace)
    return value, command_start, end


def _optional_command_argument(source: str, command: str, start: int = 0) -> tuple[str, int, int] | None:
    try:
        return _command_argument(source, command, start)
    except ValueError:
        return None


def _strip_comments(source: str) -> str:
    cleaned: list[str] = []
    for line in source.splitlines():
        match = re.search(r"(?<!\\)%", line)
        cleaned.append(line[:match.start()] if match else line)
    return "\n".join(cleaned)


def _extract_macros(source: str) -> dict[str, str]:
    macros: dict[str, str] = {}
    cursor = 0
    marker = re.compile(r"\\newcommand\s*\{\\([A-Za-z@]+)\}\s*\{")
    while match := marker.search(source, cursor):
        value, end = _extract_group(source, match.end() - 1)
        macros[match.group(1)] = value.strip()
        cursor = end
    return macros


def _expand_macros(source: str, macros: dict[str, str]) -> str:
    result = source
    for _ in range(12):
        previous = result
        for name in sorted(macros, key=len, reverse=True):
            result = re.sub(rf"\\{re.escape(name)}(?:\{{\}})?", lambda _: macros[name], result)
        if result == previous:
            break
    return result


def _collapse_source(source: str) -> str:
    source = _strip_comments(source)
    return re.sub(r"\s+", " ", source).strip()


def _plain_latex(source: str) -> str:
    """Return readable text for placeholder descriptions and diagnostics."""
    text = source.replace("---", "—").replace(r"\slash{}", "/")
    text = text.replace(r"\pm", "±").replace(r"\%", "%")
    for _ in range(8):
        previous = text
        text = re.sub(r"\\(?:textbf|textit|placeholder)\{([^{}]*)\}", r"\1", text)
        if text == previous:
            break
    text = text.replace("$", "").replace("~", " ")
    text = re.sub(r"\\[A-Za-z@]+(?:\{\})?", "", text)
    return _collapse_source(text.replace("{", "").replace("}", ""))


def _paragraph_fragments(source: str) -> list[str]:
    source = _strip_comments(source).strip()
    if not source:
        return []
    return [_collapse_source(part) for part in re.split(r"\n\s*\n", source) if part.strip()]


def _set_run_font(run, name: str = BODY_FONT, size: float | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)


def _append_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "555555")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    properties.append(fonts)
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _append_inline(
    paragraph,
    source: str,
    *,
    bold: bool = False,
    italic: bool = False,
    refs: dict[str, str] | None = None,
) -> None:
    refs = refs or {"tab:componentes": "1", "fig:indice": "1", "fig:componentes": "2"}
    source = source.replace("---", "—")
    index = 0
    buffer: list[str] = []

    def flush() -> None:
        if not buffer:
            return
        text = "".join(buffer)
        buffer.clear()
        if text:
            run = paragraph.add_run(text)
            _set_run_font(run)
            run.bold = bold
            run.italic = italic

    while index < len(source):
        if source[index] == "$":
            closing = source.find("$", index + 1)
            if closing != -1:
                math = source[index + 1:closing].replace("\\pm", "±")
                buffer.append(math)
                index = closing + 1
                continue
        if source[index] == "\\":
            command_match = re.match(r"\\([A-Za-z@]+)", source[index:])
            if command_match:
                command = command_match.group(1)
                command_end = index + len(command_match.group(0))
                while command_end < len(source) and source[command_end].isspace():
                    command_end += 1
                if command in {"textbf", "textit", "placeholder", "url", "ref"} and command_end < len(source) and source[command_end] == "{":
                    flush()
                    value, end = _extract_group(source, command_end)
                    if command == "textbf":
                        _append_inline(paragraph, value, bold=True or bold, italic=italic, refs=refs)
                    elif command == "textit":
                        _append_inline(paragraph, value, bold=bold, italic=True, refs=refs)
                    elif command == "placeholder":
                        run = paragraph.add_run(f"[{_plain_latex(value)}]")
                        _set_run_font(run)
                        run.bold = True
                        run.font.color.rgb = PLACEHOLDER_RED
                    elif command == "url":
                        _append_hyperlink(paragraph, value, value)
                    else:
                        run = paragraph.add_run(refs.get(value, value))
                        _set_run_font(run)
                        run.bold = bold
                        run.italic = italic
                    index = end
                    continue
                if command == "slash":
                    buffer.append("/")
                    index = command_end + (2 if source[command_end:command_end + 2] == "{}" else 0)
                    continue
                if command in {"par", "@"}:
                    index = command_end
                    continue
                # Unknown presentation commands are discarded, but their braced
                # text remains available to the normal parser.
                index = command_end
                continue
            if index + 1 < len(source) and source[index + 1] in "%&#_$":
                buffer.append(source[index + 1])
                index += 2
                continue
        char = source[index]
        if char == "~":
            buffer.append("\u00a0")
        elif char not in "{}":
            buffer.append(char)
        index += 1
    flush()


def _set_paragraph_body(paragraph, *, first_line: bool = True) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    # An exact baseline reproduces LuaLaTeX's one-and-a-half spacing more
    # consistently than Word's renderer-dependent "1.5 lines" setting.
    fmt.line_spacing = Pt(15)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Mm(5.3) if first_line else Mm(0)
    fmt.widow_control = True


def _add_body_paragraph(document, source: str, *, first_line: bool = True):
    paragraph = document.add_paragraph(style="Normal")
    _set_paragraph_body(paragraph, first_line=first_line)
    _append_inline(paragraph, source)
    return paragraph


def _set_cell_margins(cell, top: int = 60, start: int = 80, bottom: int = 60, end: int = 80) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: Iterable[int]) -> None:
    widths = list(widths)
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    # Align the visible table edge with surrounding paragraph text after the
    # start cell inset, per Word's table-positioning model.
    indent.set(qn("w:w"), "80")
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for column_width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(column_width))
        grid.append(column)

    for row in table.rows:
        no_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(no_split)
        for cell, column_width in zip(row.cells, widths):
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:w"), str(column_width))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_borders(cell, **borders: dict[str, str]) -> None:
    properties = cell._tc.get_or_add_tcPr()
    border_group = properties.first_child_found_in("w:tcBorders")
    if border_group is None:
        border_group = OxmlElement("w:tcBorders")
        properties.append(border_group)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        edge_element = border_group.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            border_group.append(edge_element)
        values = borders.get(edge, {"val": "nil"})
        for key, value in values.items():
            edge_element.set(qn(f"w:{key}"), value)


def _next_numbering_id(numbering, tag: str, attribute: str) -> int:
    values = [int(node.get(qn(attribute))) for node in numbering.findall(qn(tag))]
    return max(values, default=0) + 1


def _add_numbering(document, *, bullet: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_id = _next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = _next_numbering_id(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(fmt)
    if not bullet:
        # Bind the decimal definition to Heading 1 so Microsoft Word treats it
        # as native section numbering rather than guessing a list marker.
        paragraph_style = OxmlElement("w:pStyle")
        paragraph_style.set(qn("w:val"), "Heading1")
        level.append(paragraph_style)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab" if bullet else "space")
    level.append(suffix)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if bullet else "%1")
    level.append(text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    ppr = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    if bullet:
        indent.set(qn("w:left"), "720")
        indent.set(qn("w:hanging"), "360")
    else:
        indent.set(qn("w:left"), "0")
        indent.set(qn("w:hanging"), "0")
    ppr.append(indent)
    level.append(ppr)
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    rpr.append(fonts)
    if not bullet:
        bold = OxmlElement("w:b")
        rpr.append(bold)
    level.append(rpr)
    abstract.append(level)

    # OOXML requires every w:abstractNum to precede every w:num. Appending a
    # new abstract definition after python-docx's built-in w:num elements is
    # tolerated by LibreOffice but can make Word substitute the wrong marker.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, num_id: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    num_properties = properties.find(qn("w:numPr"))
    if num_properties is None:
        num_properties = OxmlElement("w:numPr")
        properties.insert(0, num_properties)
    else:
        for child in list(num_properties):
            num_properties.remove(child)
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    identifier = OxmlElement("w:numId")
    identifier.set(qn("w:val"), str(num_id))
    num_properties.append(level)
    num_properties.append(identifier)


def _set_paragraph_bottom_border(paragraph, color: str = "BFBFBF", size: str = "4") -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _configure_styles(document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.paragraph_format.line_spacing = Pt(15)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    heading = styles["Heading 1"]
    heading.font.name = BODY_FONT
    heading.font.size = Pt(16)
    heading.font.bold = True
    heading.font.color.rgb = BLACK
    heading._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(9)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.keep_together = True

    for style_name, size, italic in (
        ("IDC Caption", 10.5, True),
        ("IDC Source", 9.5, True),
        ("IDC Header", 9.0, True),
    ):
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[style_name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.italic = italic
        style.font.color.rgb = BLACK
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4 if style_name == "IDC Caption" else 8)
        style.paragraph_format.keep_with_next = style_name == "IDC Caption"


def _configure_page(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(30)
    section.right_margin = Mm(30)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)


def _clear_container(container) -> None:
    for paragraph in list(container.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)
    for table in list(container.tables):
        table._element.getparent().remove(table._element)


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, display, end])
    _set_run_font(run, size=10)


def _restart_page_number(section, start: int) -> None:
    properties = section._sectPr
    page_number = properties.find(qn("w:pgNumType"))
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        properties.append(page_number)
    page_number.set(qn("w:start"), str(start))


def _configure_content_header_footer(section, title: str) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    _clear_container(section.header)
    _clear_container(section.footer)
    # A right-aligned tab gives us the two-ended running header without a
    # layout table, which keeps the document's accessibility tree clean.
    header = section.header.add_paragraph(style="IDC Header")
    header.paragraph_format.tab_stops.add_tab_stop(Mm(150), WD_TAB_ALIGNMENT.RIGHT)
    _append_inline(header, title, italic=True)
    header.add_run("\t")
    _append_inline(header, "FGVcemif / FGV-EAESP", italic=True)
    _set_paragraph_bottom_border(header, color="BFBFBF", size="4")
    _add_page_field(section.footer.add_paragraph())


def _add_cover(document, macros: dict[str, str], logo_path: Path) -> None:
    section = document.sections[0]
    _configure_page(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    _clear_container(section.header)
    _clear_container(section.footer)

    logo = document.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo.paragraph_format.line_spacing = 1.0
    logo.paragraph_format.space_after = Mm(18)
    logo_drawing = logo.add_run().add_picture(str(logo_path), width=Mm(79))
    logo_drawing._inline.docPr.set("descr", "Logotipo FGV EAESP e FGVcemif")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Mm(12)
    title.paragraph_format.space_after = Mm(3)
    title_run = title.add_run(_collapse_source(macros["reporttitle"]))
    _set_run_font(title_run, BODY_FONT, 17.3)
    title_run.bold = True

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Mm(12)
    _append_inline(subtitle, _expand_macros(macros["reportsubtitle"], macros), italic=True)
    for run in subtitle.runs:
        _set_run_font(run, BODY_FONT, 12)

    authors = document.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    authors.paragraph_format.line_spacing = 1.15
    for position, name in enumerate((macros["authorone"], macros["authortwo"], macros["authorthree"])):
        if position:
            authors.add_run().add_break()
        run = authors.add_run(_collapse_source(name))
        _set_run_font(run, BODY_FONT, 12)

    date = section.footer.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _append_inline(date, _expand_macros(macros["reportdate"], macros))
    for run in date.runs:
        _set_run_font(run, BODY_FONT, 11)


def _section_heading(document, title: str, heading_num_id: int, *, numbered: bool = True):
    paragraph = document.add_paragraph(style="Heading 1")
    if numbered:
        _apply_numbering(paragraph, heading_num_id)
    _append_inline(paragraph, title)
    return paragraph


def _caption(document, label: str, number: int, source: str):
    paragraph = document.add_paragraph(style="IDC Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"{label} {number}. ")
    _set_run_font(run, BODY_FONT, 10.5)
    run.bold = True
    _append_inline(paragraph, source, italic=True)
    return paragraph


def _source_note(document, source: str):
    paragraph = document.add_paragraph(style="IDC Source")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = paragraph.add_run("Fonte: ")
    _set_run_font(prefix, BODY_FONT, 9.5)
    prefix.italic = True
    _append_inline(paragraph, source, italic=True)
    return paragraph


def _parse_table(table_block: str) -> tuple[str, list[list[str]], str | None]:
    caption = _command_argument(table_block, "caption")[0]
    source_match = _optional_command_argument(table_block, "fonte")
    source = source_match[0] if source_match else None
    tabular_start = re.search(r"\\begin\{tabular\}\s*", table_block)
    tabular_end = table_block.find(r"\end{tabular}")
    if not tabular_start or tabular_end == -1:
        raise ValueError("IDC results table is missing its tabular environment")
    column_brace = table_block.find("{", tabular_start.end())
    if column_brace == -1:
        raise ValueError("IDC tabular environment is missing its column specification")
    _, tabular_content_start = _extract_group(table_block, column_brace)
    tabular = table_block[tabular_content_start:tabular_end]
    tabular = re.sub(r"\\(?:toprule|midrule|bottomrule)", "", tabular)
    rows: list[list[str]] = []
    for row_source in re.split(r"\\\\", tabular):
        cells = [_collapse_source(cell) for cell in row_source.split("&")]
        if len(cells) == 3 and any(cells):
            rows.append(cells)
    return caption, rows, source


def _extract_figure(figure_block: str, number: int) -> dict[str, str | int]:
    file_match = _optional_command_argument(figure_block, "IfFileExists")
    if file_match:
        image_name = _collapse_source(file_match[0])
    else:
        include = re.search(r"\\includegraphics(?:\[[^]]*\])?\{([^}]+)\}", figure_block)
        if not include:
            raise ValueError("Figure has no image filename")
        image_name = include.group(1)
    return {
        "number": number,
        "image": image_name,
        "caption": _command_argument(figure_block, "caption")[0],
        "source": _command_argument(figure_block, "fonte")[0],
    }


def _parse_report(source: str) -> dict:
    macros = _extract_macros(source)
    expanded = _expand_macros(source, macros)

    summary_marker = expanded.find(r"\section{Resumo}")
    results_marker = expanded.find(r"\section{Resultados de", summary_marker)
    trajectory_marker = expanded.find(r"\section{Trajetória do índice}", results_marker)
    next_marker = expanded.find(r"\section{Próxima atualização}", trajectory_marker)
    notes_marker = expanded.find(r"\section*{Notas}", next_marker)
    annex_marker = expanded.find(r"\section*{Anexo de figuras}", notes_marker)
    if min(summary_marker, results_marker, trajectory_marker, next_marker, notes_marker, annex_marker) < 0:
        raise ValueError("The LaTeX report does not match the IDC section structure")

    summary_end = summary_marker + len(r"\section{Resumo}")
    results_title, _, results_content_start = _command_argument(expanded, "section", results_marker)
    summary = _paragraph_fragments(expanded[summary_end:results_marker])

    table_start = expanded.find(r"\begin{table}", results_content_start)
    table_end = expanded.find(r"\end{table}", table_start)
    if table_start == -1 or table_end == -1:
        raise ValueError("The IDC results table was not found")
    table_end += len(r"\end{table}")
    table_caption, table_rows, table_source = _parse_table(expanded[table_start:table_end])

    after_table = table_end
    external_source = _optional_command_argument(expanded, "fonte", table_end)
    if external_source and external_source[1] < trajectory_marker:
        table_source = external_source[0]
        after_table = external_source[2]
    if table_source is None:
        raise ValueError("The IDC results table is missing its source note")

    item_start = expanded.find(r"\begin{itemize}", after_table, trajectory_marker)
    item_end = expanded.find(r"\end{itemize}", item_start, trajectory_marker)
    if item_start == -1 or item_end == -1:
        raise ValueError("The IDC component list was not found")
    results_intro = _paragraph_fragments(expanded[after_table:item_start])
    item_block = expanded[item_start + len(r"\begin{itemize}"):item_end]
    items = [_collapse_source(item) for item in re.split(r"\\item\s+", item_block) if item.strip()]
    results_summary = _paragraph_fragments(expanded[item_end + len(r"\end{itemize}"):trajectory_marker])

    trajectory_start = trajectory_marker + len(r"\section{Trajetória do índice}")
    trajectory_intro = _paragraph_fragments(expanded[trajectory_start:next_marker])

    next_title, _, next_start = _command_argument(expanded, "section", next_marker)
    next_update = _paragraph_fragments(expanded[next_start:notes_marker])
    notes_start = notes_marker + len(r"\section*{Notas}")
    notes = _paragraph_fragments(expanded[notes_start:annex_marker])

    annex_start = annex_marker + len(r"\section*{Anexo de figuras}")
    document_end = expanded.find(r"\end{document}", annex_start)
    figure_matches = list(re.finditer(
        r"\\begin\{figure\}(?:\[[^]]*\])?.*?\\end\{figure\}",
        expanded[annex_start:document_end],
        re.S,
    ))
    if len(figure_matches) != 2:
        raise ValueError(f"Expected two IDC annex figures, found {len(figure_matches)}")
    figures = [_extract_figure(match.group(0), number) for number, match in enumerate(figure_matches, 1)]

    return {
        "macros": {name: _expand_macros(value, macros) for name, value in macros.items()},
        "summary": summary,
        "results_title": results_title,
        "table_caption": table_caption,
        "table_rows": table_rows,
        "table_source": table_source,
        "results_intro": results_intro,
        "items": items,
        "results_summary": results_summary,
        "trajectory_intro": trajectory_intro,
        "figures": figures,
        "next_title": next_title,
        "next_update": next_update,
        "notes": notes,
        "annex_title": "Anexo de figuras",
    }


def _add_results_table(document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=3)
    table.autofit = False
    widths = [4535, 1700, 2269]  # 80 mm, 30 mm, 40 mm.
    _set_table_geometry(table, widths)
    for row_index, row in enumerate(rows):
        for column_index, source in enumerate(row):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            _append_inline(paragraph, source)
            for run in paragraph.runs:
                _set_run_font(run, BODY_FONT, 10.5)
            border_values: dict[str, dict[str, str]] = {}
            if row_index == 0:
                border_values["top"] = {"val": "single", "sz": "8", "color": "000000"}
                border_values["bottom"] = {"val": "single", "sz": "4", "color": "000000"}
            if row_index == len(rows) - 1:
                border_values["bottom"] = {"val": "single", "sz": "8", "color": "000000"}
            _set_cell_borders(cell, **border_values)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))


def _add_figure(document, figure: dict, assets_dir: Path) -> None:
    image_path = assets_dir / str(figure["image"])
    if image_path.exists():
        image = DocxImage.from_file(str(image_path))
        scaled_height_mm = FIGURE_WIDTH_MM * image.px_height / image.px_width
        maximum_height_mm = FIGURE_MAX_HEIGHT_MM[int(figure["number"])]
        if scaled_height_mm > maximum_height_mm:
            raise ValueError(
                f"{image_path.name} would be {scaled_height_mm:.1f} mm tall at full text width; "
                f"maximum for figure {figure['number']} is {maximum_height_mm:.1f} mm. "
                "Correct the source chart aspect ratio instead of shrinking the figure."
            )
        image_paragraph = document.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Inline drawings must use an auto-expanding line box. Inheriting the
        # body's exact baseline clips tall charts in LibreOffice and Word.
        image_paragraph.paragraph_format.line_spacing = 1.0
        image_paragraph.paragraph_format.space_after = Pt(4)
        image_paragraph.paragraph_format.keep_together = True
        image_paragraph.paragraph_format.keep_with_next = True
        run = image_paragraph.add_run()
        drawing = run.add_picture(str(image_path), width=Mm(FIGURE_WIDTH_MM))
        doc_pr = drawing._inline.docPr
        doc_pr.set("descr", f"{figure['caption']}")
    else:
        # The LaTeX template uses a large framed placeholder on each annex
        # page. A one-cell Word table mirrors the box while remaining editable.
        placeholder_table = document.add_table(rows=1, cols=1)
        placeholder_table.autofit = False
        _set_table_geometry(placeholder_table, [8504])  # 150 mm content width.
        placeholder_row = placeholder_table.rows[0]
        placeholder_row.height = Mm(FIGURE_PLACEHOLDER_HEIGHT_MM[int(figure["number"])])
        placeholder_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        placeholder_cell = placeholder_table.cell(0, 0)
        placeholder_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_borders(
            placeholder_cell,
            top={"val": "single", "sz": "6", "color": "A0A0A0"},
            left={"val": "single", "sz": "6", "color": "A0A0A0"},
            bottom={"val": "single", "sz": "6", "color": "A0A0A0"},
            right={"val": "single", "sz": "6", "color": "A0A0A0"},
        )
        image_paragraph = placeholder_cell.paragraphs[0]
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.line_spacing = 1.0
        image_paragraph.paragraph_format.space_before = Pt(0)
        image_paragraph.paragraph_format.space_after = Pt(0)
        placeholder = image_paragraph.add_run(f"[{figure['image']}]")
        _set_run_font(placeholder, SANS_FONT, 10)
        placeholder.font.color.rgb = GRAY
    caption = _caption(document, "Figura", int(figure["number"]), str(figure["caption"]))
    caption.paragraph_format.keep_together = True
    source = _source_note(document, str(figure["source"]))
    source.paragraph_format.keep_together = True


def build_docx(tex_path: Path, output_path: Path, *, assets_dir: Path | None = None, require_filled: bool = False) -> None:
    source = tex_path.read_text(encoding="utf-8")
    if require_filled and r"\placeholder{" in source:
        raise ValueError("Filled monthly reports cannot contain \\placeholder{...} commands")
    report = _parse_report(source)
    macros = report["macros"]
    assets_dir = assets_dir or tex_path.parent
    logo_path = assets_dir / "logo.png"
    if not logo_path.exists():
        raise FileNotFoundError(f"Logo not found: {logo_path}")
    figure_paths = [assets_dir / str(figure["image"]) for figure in report["figures"]]
    if require_filled:
        missing_figures = [path.name for path in figure_paths if not path.exists()]
        if missing_figures:
            raise FileNotFoundError(f"Filled monthly report is missing figure assets: {', '.join(missing_figures)}")
    document = Document()
    _configure_styles(document)
    heading_num_id = _add_numbering(document, bullet=False)
    bullet_num_id = _add_numbering(document, bullet=True)
    document.core_properties.title = _collapse_source(macros["reporttitle"])
    document.core_properties.subject = "Nota Técnica de Atualização do IDC"
    document.core_properties.author = "; ".join(
        _collapse_source(macros[name]) for name in ("authorone", "authortwo", "authorthree")
    )
    document.core_properties.keywords = "crédito, endividamento, Brasil, BCB, FGV"

    _add_cover(document, macros, logo_path)

    content = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_page(content)
    _restart_page_number(content, 1)
    content.different_first_page_header_footer = True
    _configure_content_header_footer(content, _collapse_source(macros["reporttitle"]))
    content.first_page_header.is_linked_to_previous = False
    content.first_page_footer.is_linked_to_previous = False
    _clear_container(content.first_page_header)
    _clear_container(content.first_page_footer)
    blank_rule = content.first_page_header.add_paragraph()
    blank_rule.paragraph_format.space_after = Pt(0)
    _set_paragraph_bottom_border(blank_rule)
    _add_page_field(content.first_page_footer.add_paragraph())

    _section_heading(document, "Resumo", heading_num_id)
    for index, paragraph in enumerate(report["summary"]):
        _add_body_paragraph(document, paragraph, first_line=index > 0)

    _section_heading(document, report["results_title"], heading_num_id)
    _caption(document, "Tabela", 1, report["table_caption"])
    _add_results_table(document, report["table_rows"])
    _source_note(document, report["table_source"])
    for paragraph in report["results_intro"]:
        _add_body_paragraph(document, paragraph)
    for item in report["items"]:
        paragraph = _add_body_paragraph(document, item, first_line=False)
        paragraph.paragraph_format.left_indent = Mm(8)
        paragraph.paragraph_format.first_line_indent = Mm(0)
        paragraph.paragraph_format.space_before = Pt(3)
        _apply_numbering(paragraph, bullet_num_id)
    for paragraph in report["results_summary"]:
        _add_body_paragraph(document, paragraph)

    _section_heading(document, "Trajetória do índice", heading_num_id)
    for index, paragraph in enumerate(report["trajectory_intro"]):
        _add_body_paragraph(document, paragraph, first_line=index > 0)
    _section_heading(document, report["next_title"], heading_num_id)
    for index, paragraph in enumerate(report["next_update"]):
        _add_body_paragraph(document, paragraph, first_line=index > 0)
    _section_heading(document, "Notas", heading_num_id, numbered=False)
    for index, paragraph in enumerate(report["notes"]):
        _add_body_paragraph(document, paragraph, first_line=index > 0)

    # The annex isolates visuals from the flowing report body. Each chart gets
    # a full-width page, so neither chart is shrunk or displaced by text.
    annex = _section_heading(document, report["annex_title"], heading_num_id, numbered=False)
    annex.paragraph_format.page_break_before = True
    _add_figure(document, report["figures"][0], assets_dir)

    document.add_page_break()
    _add_figure(document, report["figures"][1], assets_dir)

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("tex", type=Path, help="Filled IDC LaTeX report")
    parser.add_argument("output", type=Path, help="Destination DOCX path")
    parser.add_argument("--assets-dir", type=Path, help="Directory containing logo.png and report figures")
    parser.add_argument(
        "--require-filled",
        action="store_true",
        help="Fail if the source still contains template placeholders",
    )
    args = parser.parse_args()
    build_docx(args.tex, args.output, assets_dir=args.assets_dir, require_filled=args.require_filled)
    print(f"DOCX generated: {args.output}")


if __name__ == "__main__":
    main()
