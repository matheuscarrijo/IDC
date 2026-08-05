import tempfile
import unittest
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.shared import Mm

from src.build_report_docx import _add_figure, _parse_report, build_docx


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_TEX = ROOT / "outputs/report/template-latex/template.tex"
TEMPLATE_ASSETS = TEMPLATE_TEX.parent
FILLED_TEX = ROOT / "outputs/report/update-202607/idc-update-202607.tex"
FILLED_ASSETS = FILLED_TEX.parent
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{W_NS}}}"


class BuildReportDocxTests(unittest.TestCase):
    def test_parser_reads_the_supported_idc_template(self):
        report = _parse_report(TEMPLATE_TEX.read_text(encoding="utf-8"))

        self.assertEqual(report["results_title"], r"Resultados de \placeholder{mês por extenso de ano}")
        self.assertEqual(len(report["table_rows"]), 5)
        self.assertEqual(report["table_rows"][0], ["Indicador", "Valor bruto", "Valor normalizado"])
        self.assertEqual([figure["image"] for figure in report["figures"]], ["index.png", "components_raw.png"])

    def test_builder_creates_an_editable_a4_document(self):
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "template.docx"
            build_docx(TEMPLATE_TEX, output, assets_dir=TEMPLATE_ASSETS)

            document = Document(output)
            self.assertEqual(len(document.sections), 2)
            self.assertAlmostEqual(document.sections[0].page_width.mm, 210, places=1)
            self.assertAlmostEqual(document.sections[0].page_height.mm, 297, places=1)
            # One data table plus two editable chart-placeholder boxes.
            self.assertEqual(len(document.tables), 3)
            self.assertIn("Indicador", document.tables[0].cell(0, 0).text)
            visible_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("Índice de Desconforto de Crédito", visible_text)
            self.assertIn("Trajetória do índice", visible_text)

    def test_heading_numbering_is_native_decimal_and_word_compatible(self):
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "template.docx"
            build_docx(TEMPLATE_TEX, output, assets_dir=TEMPLATE_ASSETS)

            with zipfile.ZipFile(output) as archive:
                numbering = ElementTree.fromstring(archive.read("word/numbering.xml"))
                document_xml = ElementTree.fromstring(archive.read("word/document.xml"))

            children = list(numbering)
            first_num_index = next(
                index for index, child in enumerate(children) if child.tag == f"{W}num"
            )
            self.assertTrue(
                all(child.tag == f"{W}abstractNum" for child in children[:first_num_index])
            )

            abstract_by_id = {
                node.get(f"{W}abstractNumId"): node
                for node in numbering.findall(f"{W}abstractNum")
            }
            abstract_id_by_num = {
                node.get(f"{W}numId"): node.find(f"{W}abstractNumId").get(f"{W}val")
                for node in numbering.findall(f"{W}num")
            }

            numbered_headings = []
            for paragraph in document_xml.findall(f".//{W}p"):
                properties = paragraph.find(f"{W}pPr")
                if properties is None:
                    continue
                style = properties.find(f"{W}pStyle")
                num_id = properties.find(f"{W}numPr/{W}numId")
                if style is not None and style.get(f"{W}val") == "Heading1" and num_id is not None:
                    numbered_headings.append(num_id.get(f"{W}val"))

            self.assertEqual(len(numbered_headings), 4)
            for num_id in numbered_headings:
                abstract = abstract_by_id[abstract_id_by_num[num_id]]
                self.assertEqual(abstract.find(f"{W}lvl/{W}numFmt").get(f"{W}val"), "decimal")
                self.assertEqual(abstract.find(f"{W}lvl/{W}pStyle").get(f"{W}val"), "Heading1")
                self.assertEqual(abstract.find(f"{W}lvl/{W}lvlText").get(f"{W}val"), "%1")

    def test_require_filled_rejects_the_template(self):
        with tempfile.TemporaryDirectory() as directory:
            with self.assertRaisesRegex(ValueError, "cannot contain"):
                build_docx(
                    TEMPLATE_TEX,
                    Path(directory) / "should-not-exist.docx",
                    assets_dir=TEMPLATE_ASSETS,
                    require_filled=True,
                )

    def test_filled_figures_are_full_width_and_kept_with_their_captions(self):
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "filled.docx"
            build_docx(FILLED_TEX, output, assets_dir=FILLED_ASSETS, require_filled=True)

            document = Document(output)
            drawing_paragraphs = [
                paragraph for paragraph in document.paragraphs
                if paragraph._p.xpath(".//w:drawing")
            ]
            # Cover logo plus the two report figures.
            self.assertEqual(len(drawing_paragraphs), 3)
            for paragraph in drawing_paragraphs[1:]:
                extent = paragraph._p.xpath(".//wp:extent")[0]
                self.assertEqual(int(extent.get("cx")), int(Mm(150)))
                self.assertTrue(paragraph.paragraph_format.keep_with_next)

            # The first chart must also use its annex page vertically; a
            # very wide source would recreate the large blank gap in the PDF.
            first_figure_extent = drawing_paragraphs[1]._p.xpath(".//wp:extent")[0]
            self.assertGreaterEqual(int(first_figure_extent.get("cy")), int(Mm(130)))

            trajectory = next(
                paragraph for paragraph in document.paragraphs
                if paragraph.text == "Trajetória do índice"
            )
            self.assertIsNone(trajectory.paragraph_format.page_break_before)

            annex = next(
                paragraph for paragraph in document.paragraphs
                if paragraph.text == "Anexo de figuras"
            )
            self.assertTrue(annex.paragraph_format.page_break_before)

    def test_builder_rejects_a_tall_chart_instead_of_shrinking_it(self):
        document = Document()
        too_tall_for_figure_one = {
            "number": 1,
            "image": "components_raw.png",
            "caption": "Teste",
            "source": "Teste.",
        }

        with self.assertRaisesRegex(ValueError, "instead of shrinking"):
            _add_figure(document, too_tall_for_figure_one, FILLED_ASSETS)


if __name__ == "__main__":
    unittest.main()
